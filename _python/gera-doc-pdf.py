import pandas
import os
from docx import Document
 
cursos = ['EA', 'EB', 'EF', 'EM', 'EP', 'EQD', 'EQN']
 
for curso in cursos: 
 
    print(f' {curso} '.center(60,'#'))
    
    df = pandas.read_json(f'{curso}.json', orient='index')

    # print(df.columns.values)

    for index, discpln in df.iterrows():

        print(f'{discpln["sigla"]} - {discpln["nome"]}')

        doc = Document()
        try:
            os.mkdir(f'../{curso}')
        except  FileExistsError:
            pass
        docname = f'../assets/cursos/{curso}/{discpln["sigla"]}.docx'

        doc.add_heading(f'{discpln["sigla"]} - {discpln["nome"]}')
        doc.add_heading(f'{discpln["nome_en"]}', level=3)
        doc.add_paragraph()

        # Dados gerais

        p = doc.add_paragraph(style = 'List Number')
        p.add_run(f'Créditos-aula: {discpln["CA"]}\n')
        p.add_run(f'Créditos-trabalho: {discpln["CT"]}\n')
        p.add_run(f'Carga horária: {discpln["CH"]}\n')
        p.add_run(f'Semestre ideal: {discpln["semestre"]}\n')
        p.add_run(f'Ativação: {discpln["ativacao"]}\n')
        p.add_run(f'Departamento: {discpln["departamento"]}')

        # Objetivos

        doc.add_heading(f'Objetivos', level=2)

        doc.add_paragraph(f'{discpln["objetivos"]}')

        if discpln["abstract"]:
            p = doc.add_paragraph()
            p.add_run(f'{discpln["objectives"]}').italic = True

        # Docentes

        doc.add_heading(f'Docente(s) Responsável(eis) ', level=2)

        profs = discpln["docentes"]
        nprofs = discpln["ndoc"]

        if nprofs:
            p = doc.add_paragraph(style='List Bullet')
            for i in range(nprofs-1):
                p.add_run(f'{profs[i]}\n')
            p.add_run(f'{profs[-1]}')

        # programa resumido

        doc.add_heading(f'Programa resumido', level=2)

        doc.add_paragraph(f'{discpln["resumo"]}')

        if discpln["abstract"]:
            p = doc.add_paragraph()
            p.add_run(f'{discpln["abstract"]}').italic = True
        # programa

        doc.add_heading(f'Programa', level=2)

        doc.add_paragraph(f'{discpln["programa"]}')

        if discpln["program"]:
            p = doc.add_paragraph()
            p.add_run(f'{discpln["program"]}').italic = True

        # avaliação

        doc.add_heading('Avaliação', level=2)

        p = doc.add_paragraph( style='List Bullet')
        p.add_run('Método: ').bold = True
        p.add_run(f'{discpln["metodo"]}\n')

        p.add_run('Critério: ').bold = True
        p.add_run(f'{discpln["criterio"]}\n')

        p.add_run('Norma de recuperação: ').bold = True
        p.add_run(f'{discpln["exame"]}')

        # bibliografia

        doc.add_heading('Bibliografia', level=2)

        doc.add_paragraph(f'{discpln["bibliografia"]}')

        # Requisitos

        nr = discpln['requisitos']
        if nr:
            doc.add_heading('Requisitos', level=2)

            p = doc.add_paragraph(style='List Bullet')
            for k, req in nr.items():
                p.add_run(f"{req['sigla']} - {req['nome']} ({req['tipo']})\n")

        # salvando

        doc.save(docname)

        # exportando pdf

        os.system(f'abiword --to=pdf {docname}')
